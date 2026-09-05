# -*- coding: utf-8 -*-
"""Создаёт рабочие docx-шаблоны с плейсхолдерами {{Key}} (без пробелов) из исходников.

Заполняются только поля, для которых есть данные в CRM; поля событий
(даты передачи/монтажа, состав работ/изделий) остаются пустыми строками
оригинала для ручного дополнения.
"""
import os, re
import docx

SRC = r"D:\CLAUDE\Project\unified-crm-finance\documents"
OUT = r"D:\CLAUDE\Project\unified-crm-finance\apps\web\src\server\templates"
os.makedirs(OUT, exist_ok=True)

def norm(s):
    return re.sub(r"\s+", " ", s).strip()

def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for sub in cell.tables:
                    for r2 in sub.rows:
                        for c2 in r2.cells:
                            for p in c2.paragraphs:
                                yield p
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            if hf and not hf.is_linked_to_previous:
                for p in hf.paragraphs:
                    yield p

def write_para(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

def apply_rules(doc, rules):
    changed = 0
    for p in iter_paragraphs(doc):
        raw = norm(p.text)
        if not raw:
            continue
        new = raw
        for rx, repl in rules:
            if isinstance(repl, str):
                new, n = re.subn(rx, repl, new)
            else:
                new, n = rx.sub(repl, new)
            changed += n
        if new != raw:
            write_para(p, new)
    return changed

def apply_rules_with_sequence(doc, rules, underscore_placeholders):
    """underscore_placeholders применяются по порядку к абзацам-строкам из подчёркиваний."""
    changed = 0
    seq = list(underscore_placeholders)
    for p in iter_paragraphs(doc):
        raw = norm(p.text)
        if not raw:
            continue
        if re.fullmatch(r"_{5,}", raw) and seq:
            write_para(p, seq.pop(0))
            changed += 1
            continue
        new = raw
        for rx, repl in rules:
            if isinstance(repl, str):
                new, n = re.subn(rx, repl, new)
            else:
                new, n = rx.sub(repl, new)
            changed += n
        if new != raw:
            write_para(p, new)
    return changed

# --- Общие блоки ---
def party_intro():
    return [(re.compile(r"(с одной стороны, и )_{10,}(, именуем)"), r"\1{{Client.fullName}}\2")]

def address_after_label():
    return [(re.compile(r"по адресу:\s*_{3,}"), "по адресу: {{Order.deliveryAddress}}")]

def requisites_person():
    return [
        (re.compile(r"^Ф\.И\.О\.:\s*_{3,}$"), "Ф.И.О.: {{Client.fullName}}"),
        (re.compile(r"Паспорт:\s*серия\s+_{3,}\s*№\s+_{3,}"),
         "Паспорт: серия {{Client.passport.series}} № {{Client.passport.number}}"),
        (re.compile(r"^Выдан:\s*_{3,}$"), "Выдан: {{Client.passport.issuedBy}}"),
        (re.compile(r"Дата выдачи:\s*_{3,}\s*Код подразделения:\s*_{3,}"),
         "Дата выдачи: {{Client.passport.issuedAt}} Код подразделения: {{Client.passport.code}}"),
        (re.compile(r"^Адрес:\s*_{3,}$"), "Адрес: {{Client.regAddress}}"),
        (re.compile(r"^Тел\.:\s*_{3,}$"), "Тел.: {{Client.phone}}"),
    ]

def date_cells(placeholder):
    return [(re.compile(r"^«___»\s*_{3,}\s*(20__|2026)\s*г\.?$"), "{{%s}}" % placeholder)]

def count_placeholders(path):
    d = docx.Document(path)
    return sum(len(re.findall(r"\{\{[^}]+\}\}", p.text)) for p in iter_paragraphs(d))

def convert(src_name, rules):
    doc = docx.Document(os.path.join(SRC, src_name))
    n = apply_rules(doc, rules)
    out = os.path.join(OUT, src_name)
    doc.save(out)
    print(f"{src_name}: замен={n}, плейсхолдеров={count_placeholders(out)}")

def convert_sequential(src_name, rules, underscore_placeholders):
    doc = docx.Document(os.path.join(SRC, src_name))
    n = apply_rules_with_sequence(doc, rules, underscore_placeholders)
    out = os.path.join(OUT, src_name)
    doc.save(out)
    print(f"{src_name}: замен={n}, плейсхолдеров={count_placeholders(out)}")

# --- 1. Договор купли-продажи ---
convert("sale_contract_2026.docx", [
    (re.compile(r"ДОГОВОР № _{3,} купли-продажи"), "ДОГОВОР № {{Doc.saleContract.number}} купли-продажи"),
    (re.compile(r"составляет\s+_{5,}\s*\(_{5,}\)\s*рублей\s+_{1,}\s*копеек"),
     "составляет {{Order.amountRoubles}} ({{Order.amountWords}}) рублей {{Order.amountKopecks}} копеек"),
    (re.compile(r"НДС не облагается / в том числе НДС \(нужное указать\)"), "{{Order.ndsLabel}}"),
    *date_cells("Doc.saleContract.date"),
    *party_intro(),
    *requisites_person(),
])

# --- 2. Договор на монтаж ---
convert("assembly_service_contract_2026.docx", [
    (re.compile(r"ДОГОВОР № _{3,} возмездного"), "ДОГОВОР № {{Doc.serviceContract.number}} возмездного"),
    (re.compile(r"Договору купли-продажи № _{3,} от «___»\s*_{3,}\s*20__\s*г\."),
     "Договору купли-продажи № {{Doc.saleContract.number}} от {{Doc.saleContract.date}}"),
    *date_cells("Doc.serviceContract.date"),
    *party_intro(),
    *address_after_label(),
    *requisites_person(),
])

# --- 3. Акт выполненных работ ---
convert("works_acceptance_act_2026.docx", [
    (re.compile(r"АКТ № _{3,} выполненных работ"), "АКТ № {{Doc.actWorks.number}} выполненных работ"),
    (re.compile(r"№ _{3,} от «___»\s*_{3,}\s*20__\s*г\."),
     "№ {{Doc.serviceContract.number}} от {{Doc.serviceContract.date}}"),
    *date_cells("Doc.actWorks.date"),
    *party_intro(),
    *address_after_label(),
])

# --- 4. Акт приёма-передачи Товара ---
convert("goods_transfer_act_2026.docx", [
    (re.compile(r"Договору купли-продажи № _{3,} от «___»\s*_{3,}\s*2026\s*г\."),
     "Договору купли-продажи № {{Doc.saleContract.number}} от {{Doc.saleContract.date}}"),
    *date_cells("Doc.actAcceptance.date"),
    *party_intro(),
    *requisites_person(),
])

# --- 5. Гарантийный талон (3 пустые строки заполняются по порядку) ---
convert_sequential("warranty_card_2026.docx", [
    (re.compile(r"ГАРАНТИЙНЫЙ ТАЛОН № _{3,}"), "ГАРАНТИЙНЫЙ ТАЛОН № {{Doc.guarantee.number}}"),
    (re.compile(r"№ Договора купли-продажи / Бланка-заказа:\s*_{3,}\s*/\s*№\s*_{3,}\s*ПМ-26"),
     "№ Договора купли-продажи / Бланка-заказа: {{Doc.saleContract.number}} / № {{Doc.blankOrder.number}}"),
    (re.compile(r"Дата заказа:\s*«___»\s*_{3,}\s*2026\s*г\."),
     "Дата заказа: {{Doc.saleContract.date}}"),
    (re.compile(r"Дата передачи Товара \(установки\):\s*«___»\s*_{3,}\s*20__\s*г\."),
     "Дата передачи Товара (установки): {{Doc.actAcceptance.date}}"),
    *date_cells("Doc.actAcceptance.date"),
], ["{{Client.fullName}}", "{{Order.deliveryAddress}}", "{{Order.productName}}"])

# --- 6. Памятка покупателю ---
convert("buyer_memo_2026.docx", [
    *date_cells("Doc.memo.signDate"),
])

# --- 7. Правила эксплуатации ---
convert("usage_rules_2026.docx", [
    (re.compile(r"^«_____»\s*_{3,}\s*2026\s*г\.?$"), "{{Doc.rules.signDate}}"),
    (re.compile(r"^ФИО:$"), "ФИО: {{Client.fullName}}"),
    (re.compile(r"^паспорт: серия номер$"), "паспорт: серия {{Client.passport.series}} номер {{Client.passport.number}}"),
    (re.compile(r"^Выдан:$"), "Выдан: {{Client.passport.issuedBy}}"),
    (re.compile(r"^Дата выдачи: Код подразделения:$"),
     "Дата выдачи: {{Client.passport.issuedAt}} Код подразделения: {{Client.passport.code}}"),
    (re.compile(r"^Адрес:$"), "Адрес: {{Client.regAddress}}"),
    (re.compile(r"^тел:$"), "тел: {{Client.phone}}"),
    (re.compile(r"^«\s*»\s*2026\s*г\.?$"), "{{Doc.rules.signDate}}"),
])

print("DONE")
